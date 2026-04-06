import docx, glob
for file in glob.glob("*.docx"):
    print(f"\n--- {file} ---")
    doc = docx.Document(file)
    for idx, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            try:
                if 'RECORD APU' in row.cells[1].text:
                    print(f"Table {idx}, Row {r}: {row.cells[1].text.strip()[:30]}")
                    cnt = 1
                    for r2 in table.rows[:r]:
                        if r2.cells[1].text.strip() and r2.cells[2].text.strip() and r2.cells[3].text.strip():
                            if any(p._p.xpath('.//w:numPr') for p in r2.cells[0].paragraphs):
                                cnt += 1
                    print(f"  My Counter logic gives ID: {cnt}")
            except Exception:
                pass
