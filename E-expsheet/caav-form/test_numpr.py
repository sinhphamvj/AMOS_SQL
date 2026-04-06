import docx
doc = docx.Document("CAAV-FSSD Form 586 [2] 2021 AMT Log book CAT A VJC only 2023.docx")
table = doc.tables[23]

counter = 1
for i, row in enumerate(table.rows):
    cell = row.cells[0]
    
    # Check if this cell contains numbering by looking at w:numPr
    has_num = False
    for p in cell.paragraphs:
        if p._p.xpath('.//w:numPr'):
            has_num = True
            break
            
    is_valid_row = bool(row.cells[1].text.strip() and row.cells[2].text.strip() and row.cells[3].text.strip())
    
    if is_valid_row:
        if has_num:
            print(f"[{counter}] {row.cells[1].text.strip().replace('\n',' ')[:30]}")
            counter += 1
        else:
            print(f"[NO NUM] {row.cells[1].text.strip().replace('\n',' ')[:30]}")
