import docx
doc = docx.Document("CAAV-FSSD Form 586 [2] 2021 AMT Log book CAT A VJC only 2023.docx")
table = doc.tables[23]
cnt = 1
for r in table.rows:
    p = r.cells[0].paragraphs[0]
    if p._p.xpath('.//w:numPr'):
        if 'RECORD APU' in r.cells[1].text:
            print(f"Total numbered rows before this: {cnt}")
            break
        cnt += 1
